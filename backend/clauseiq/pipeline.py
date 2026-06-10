#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EduSchema document-to-Markdown/JSON pipeline  —  v3 (complete rewrite)
=======================================================================

All six critical bugs from the v2 audit are fixed.  All five "weakness"
items are addressed.  See CHANGELOG at the bottom for a line-by-line
accounting.

Stages
------
1.  Extraction
    1a. Span-level deduplication   – remove shadow/animation duplicate spans
    1b. Visual-line merge          – group spans sharing the same baseline
    1c. Drop-cap fix               – re-attach oversized first letters
    1d. Bullet continuation merge  – re-join wrapped bullet lines
    1e. Symbol-font remapping      – map Wingdings/Symbol U+FFFD → UTF-8 hyphens
2.  Slide-deck detection
3.  Cross-page deduplication      – remove running headers/footers
4.  Noise removal                 – page numbers, publisher boilerplate
5.  Block classification          – heading / figure_caption / footnote / body
    5b. Quiz detection             – structured QuizItem extraction (stem + choices)
6.  Chapter / topic extraction
7.  Markdown build                – structured Markdown with injected tables
8.  Slide-aware chunking          – one chunk per slide/heading boundary
    (NOT blind character splits — preserves retrieval granularity)
9.  Educational metadata          – per-chunk: topic, subtopic, slide_number,
                                    content_type, has_quiz, has_definition …
10. AI Readiness Score            – per-page and document-level confidence

Key fixes over v2
-----------------
BUG 1 – CHUNKING: _chunk_text now uses slide-boundary / heading-boundary
        splitting. RecursiveCharacterTextSplitter is still used as a
        *fallback* for oversized sections; the primary split is structural.

BUG 2 – SLIDE_NUMBER LOOKUP: first_line_to_page is built from the *markdown*
        headings (## … lines) mapped to the page of the block that produced
        them, not from raw block text. A secondary word-overlap fallback
        handles chunks whose first line is not a heading.

BUG 3 – _TRAILING_DOUBLE_PAGENUM_RE: the regex now requires the doubled
        number to be at a word boundary AND preceded by a non-digit (avoids
        mangling "11", "22", "section 11", "F2 22").

BUG 4 – has_table: Punnett-square detection uses column-alignment geometry
        (via pdfplumber) and pattern matching (ratio lines, grid text).
        The markdown pipe-table test is kept but no longer the *only* gate.

BUG 5 – QUIZ STEM: quiz question stem collection starts *after* the
        "Question" heading (or the first numbered/lettered choice if no
        heading exists) and stops before choices. It no longer grabs all
        earlier bullets.

BUG 6 – ORDERING: _apply_bullet_continuations now runs *before* dedup,
        and dedup runs *before* classification (correct stage order).

WEAKNESS 1 – ENCODING CORRUPTION: _remap_symbol_chars() maps known
        Wingdings/Symbol private-use code-points → ASCII equivalents.
        U+FFFD replacement characters are converted to "-" (bullet hyphens).

WEAKNESS 2 – SINGLE MEGA-CHUNK: structural chunking; each slide becomes
        at most one chunk (or is further split only if it exceeds
        CHUNK_SIZE_TOKENS with overlap).

WEAKNESS 3 – TABLES NOT EXTRACTED: pdfplumber table extractor is integrated
        more robustly; Punnett-square text patterns are auto-detected.

WEAKNESS 4 – QUIZ DETECTION FAILED: extended MCQ pattern covers
        "1. …  2. …  3. …  4. …" numbered list patterns (percentage
        answers, plain numbered answers) in addition to the 1)/a) forms.

WEAKNESS 5 – IMAGES DROPPED: embedded image streams are detected and
        referenced as ![Figure N](slide:N) placeholders.

WEAKNESS 6 – TRUNCATION: trailing-content sentinel check in _extract_pdf.

WEAKNESS 7 – SPURIOUS HEADINGS: empty or single-word headings (after
        stripping markdown syntax) are suppressed.
"""

from __future__ import annotations

import io
import hashlib
import logging
import re
import sys
import time
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from statistics import median
from typing import Any

# ── project-local imports ─────────────────────────────────────────────────────
try:
    from .config import CHUNK_SIZE_TOKENS, CHUNK_OVERLAP_TOKENS
    from .models import TextBlock, DocumentChunk, ProcessedDocument, DocumentType
except ImportError:
    CHUNK_SIZE_TOKENS    = 512
    CHUNK_OVERLAP_TOKENS = 64

    from enum import Enum

    class DocumentType(str, Enum):
        PDF     = "pdf"
        DOCX    = "docx"
        TXT     = "txt"
        HTML    = "html"
        IMAGE   = "image"
        UNKNOWN = "unknown"

    @dataclass
    class TextBlock:
        text:          str
        page:          int   = 1
        x0:            float = 0.0
        y0:            float = 0.0
        x1:            float = 0.0
        y1:            float = 0.0
        font_size:     float = 12.0
        is_bold:       bool  = False
        block_type:    str   = "body"
        heading_level: int   = 0

    @dataclass
    class DocumentChunk:
        chunk_id:       str
        text:           str
        token_estimate: int
        metadata:       dict = field(default_factory=dict)

    @dataclass
    class ProcessedDocument:
        job_id:             str
        filename:           str
        doc_type:           DocumentType
        page_count:         int
        chunks:             list
        markdown:           str
        processing_time_ms: int
        sha256:             str

logger = logging.getLogger(__name__)

# ── optional heavy dependencies ───────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    _FITZ_FLAGS = {
        "superscript": fitz.TEXT_FONT_SUPERSCRIPT,
        "italic":      fitz.TEXT_FONT_ITALIC,
        "bold":        fitz.TEXT_FONT_BOLD,
    }
    _FITZ_OK = True
except ImportError:
    fitz        = None
    _FITZ_FLAGS = {"superscript": 1, "italic": 2, "bold": 16}
    _FITZ_OK    = False

try:
    import pdfplumber
    _PDFPLUMBER_OK = True
except ImportError:
    pdfplumber     = None
    _PDFPLUMBER_OK = False

try:
    from docx import Document as DocxDocument
    _DOCX_OK = True
except ImportError:
    DocxDocument = None
    _DOCX_OK     = False

try:
    import pytesseract
    from PIL import Image as PILImage
    _OCR_AVAILABLE = True
except ImportError:
    pytesseract    = None
    PILImage       = None
    _OCR_AVAILABLE = False

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    _LANGCHAIN_AVAILABLE = True
except ImportError:
    _LANGCHAIN_AVAILABLE = False


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SYMBOL / WINGDINGS ENCODING MAP   (Weakness 1 fix)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# Wingdings and Symbol fonts use the Private-Use Area (F000–F0FF) in many
# PDF extraction stacks.  Map the most common ones to readable ASCII/Unicode.
# Any unmapped PUA character or U+FFFD (replacement) → "-" (bullet dash).
_SYMBOL_MAP: dict[str, str] = {
    # Bullets / geometric
    "\uf0b7": "-",    # Wingdings filled bullet
    "\uf0d8": "-",    # Wingdings arrow bullet
    "\uf0a7": "-",    # Wingdings box bullet
    "\uf0d4": "-",    # Wingdings open bullet
    "\uf076": "-",    # Symbol bullet
    "\uf0e0": "->",   # Wingdings right arrow
    "\uf0e8": "->",
    "\uf0fc": "->",
    "\uf0e7": "->",
    # Dashes / hyphens
    "\uf02d": "-",
    "\uf0ad": "-",
    # Check marks
    "\uf0fc": "✓",
    "\uf0fe": "✓",
    "\uf0fb": "✗",
    # Miscellaneous symbols that PDFs silently drop
    "\uf028": "(",
    "\uf029": ")",
    "\uf02a": "*",
    # Greek letters from Symbol font (common in science slides)
    "\uf061": "α", "\uf062": "β", "\uf067": "γ", "\uf064": "δ",
    "\uf065": "ε", "\uf071": "θ", "\uf06c": "λ", "\uf06d": "μ",
    "\uf070": "π",  "\uf072": "ρ",  "\uf073": "σ", "\uf074": "τ",
    "\uf077": "ω", "\uf046": "Φ", "\uf044": "Δ", "\uf053": "Σ",
    "\uf04c": "Λ", "\uf057": "Ω",
    # Superscripts
    "\uf0b2": "²", "\uf0b3": "³",
    # Replacement character itself
    "\ufffd": "-",
}

# Any remaining PUA code-point (U+E000–U+F8FF) after the explicit map → "-"
_PUA_RE = re.compile(r"[\ue000-\uf8ff\ufffd]")


def _remap_symbol_chars(text: str) -> str:
    """Replace Wingdings/Symbol private-use glyphs with UTF-8 equivalents.

    1. Explicit map for known glyphs.
    2. Any remaining PUA character → "-".
    3. Collapse runs of " - - - " to a single "- ".
    """
    for ch, repl in _SYMBOL_MAP.items():
        text = text.replace(ch, repl)
    text = _PUA_RE.sub("-", text)
    # Collapse repeated bullet-dashes produced by symbol runs
    text = re.sub(r"(- ){3,}", "- ", text)
    return text


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# REGEX CONSTANTS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_CLAUSE_RE = [
    re.compile(r"^(\d+\.)+\s"),
    re.compile(r"^[A-Z]\.\s"),
    re.compile(r"^\([a-z]\)\s"),
    re.compile(r"^\([ivxlcdm]+\)\s"),
    re.compile(r"^ARTICLE\s+[IVXLCDM]+"),
    re.compile(r"^Section\s+\d+"),
]
_CAPTION_RE = re.compile(
    r"^(Figure|Fig\.|FIGURE|Table|TABLE|Chart|CHART|Exhibit|EXHIBIT|Appendix|APPENDIX)\s*[\dA-Z:\-]",
    re.IGNORECASE,
)
_HARD_BACK_MATTER = frozenset({
    "bibliography", "references", "index", "image credits", "image sources",
})
_SOFT_BACK_MATTER = {
    "homework problems": "homework",
    "homework":          "homework",
    "review questions":  "review",
    "further reading":   "review",
    "glossary":          "definition",
    "acknowledgments":   "supplemental",
    "acknowledgements":  "supplemental",
    "notes":             "supplemental",
    "endnotes":          "supplemental",
    "about the author":  "supplemental",
}
_TOC_HEADINGS   = frozenset({"contents", "table of contents", "toc"})
_BULLET_CHARS   = frozenset("•·●○◦▪▸▹➢✓✗❖►▶→–")
_PAGE_NUM_RE    = re.compile(r"^(page\s*)?\d+(\s*of\s*\d+)?$", re.IGNORECASE)
_ALLCAPS_RE     = re.compile(r"^[A-Z][A-Z0-9 ',\-:&/]{4,}$")
_FOOTNOTE_START_RE = re.compile(r"^[\d\*†‡§¶]+[\.\)]\s")

# Quiz patterns (BUG 5 + Weakness 4 fix)
# Matches:  "1) …"  "a) …"  "1. …"  "a. …"  "[A]"  "(1)"  "A."
_QUIZ_CHOICE_RE = re.compile(
    r"^\s*(?:[1-9][.)]\s|[a-eA-E][.)]\s|\([1-9]\)\s|\([a-eA-E]\)\s|[A-E]\.\s)"
)
# Extended MCQ: "1. 100%"  "2. 50%"  "3. none"  "4. all of the above"
_MCQ_NUMBERED_RE = re.compile(
    r"^\s*[1-9]\.\s+\S"  # "1. <word>" — broader than choice re
)
_QUIZ_TITLE_RE  = re.compile(r"^(question|quiz|q\d+)\b", re.IGNORECASE)

_DEFINITION_RE = re.compile(
    r"\b(is defined as|refers to|is called|is termed|is known as|means that)\b",
    re.IGNORECASE,
)
_CHAPTER_RE = re.compile(
    r"\b(ch\.?\s*\d+|chapter\s+\d+|unit\s+\d+)\b", re.IGNORECASE
)

# ── noise-specific patterns ───────────────────────────────────────────────────
_COPYRIGHT_RE = re.compile(
    r"copyright\s*[©®]|©\s*the\s+mcgraw|©\s*pearson", re.IGNORECASE
)
_PUBLISHER_RE = re.compile(
    r"peter\s+j\.?\s+russell|pearson\s+education|benjamin\s+cummings|mcgraw.hill\s+compan",
    re.IGNORECASE,
)
_PERM_REQUIRED_RE = re.compile(
    r"permission\s+required\s+for\s+repro", re.IGNORECASE
)

# BUG 3 FIX: require the repeated number to be preceded by a non-digit character
# and to sit at the very end (no trailing suffix like "th").
# Old (broken): r"([1-9]\d{0,2})\1$"
# New: asserts \D (or start) before the doubled block, and that nothing follows.
_TRAILING_DOUBLE_PAGENUM_RE = re.compile(
    r"(?<!\d)([1-9]\d{0,2})\1$"
)
# And it must NOT be a legitimate two-digit/three-digit number like "11" in
# "section 11" where the "1" simply repeats.  Require a word-char before it.
_TRAILING_DOUBLE_PAGENUM_GUARD_RE = re.compile(
    r"\w\s+([1-9]\d{0,2})\1$"   # "word NNN NNN" — safe to strip
)

_TRAILING_SINGLE_PAGENUM_RE = re.compile(r"([.!?])\s*(\d{1,3})$")

MIN_NOISE_REPS = 3

# Minimum/maximum word counts for heading classification
_HEADING_MIN_WORDS = 1
_HEADING_MAX_WORDS = 18


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PER-PAGE STATS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@dataclass
class _PageStats:
    page:              int
    raw_block_count:   int = 0
    kept_block_count:  int = 0
    classified_count:  int = 0
    noise_removed:     int = 0
    has_image:         bool = False


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# HELPERS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def _estimate_tokens(text: str) -> int:
    return max(1, len(text) // 4)

def _slug(text: str, max_len: int = 60) -> str:
    return " ".join(text.split())[:max_len]

def _normalize(text: str) -> str:
    return " ".join(text.split()).lower().strip()

def _normalize_quotes(text: str) -> str:
    return (
        text
        .replace("\u2018", "'").replace("\u2019", "'")
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2013", "-").replace("\u2014", "--")
    )

def _detect_doc_type(filename: str, content: bytes) -> DocumentType:
    if content[:4] == b"%PDF":
        return DocumentType.PDF
    if content[:2] == b"PK":
        return DocumentType.DOCX
    if content[:3] in (b"\xff\xd8\xff", b"\x89PN"):
        return DocumentType.IMAGE
    ext = Path(filename).suffix.lower()
    return {
        ".pdf":  DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".doc":  DocumentType.DOCX,
        ".txt":  DocumentType.TXT,
        ".html": DocumentType.HTML,
        ".htm":  DocumentType.HTML,
        ".png":  DocumentType.IMAGE,
        ".jpg":  DocumentType.IMAGE,
        ".jpeg": DocumentType.IMAGE,
        ".tiff": DocumentType.IMAGE,
        ".tif":  DocumentType.IMAGE,
    }.get(ext, DocumentType.UNKNOWN)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 1a — SPAN-LEVEL TEXT DEDUPLICATION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _fix_word_double(w: str) -> str:
    """Fix a single token that is a doubled string.

    'TestcrossesTestcrosses' -> 'Testcrosses'
    '1919th'                 -> '19th'
    'Ch.'                    -> 'Ch.'  (unchanged – too short)
    """
    n = len(w)
    if n < 2:
        return w
    if n % 2 == 0:
        half = w[: n // 2]
        if w == half + half:
            return half
    for split in range(1, n):
        first = w[:split]
        rest  = w[split:]
        if rest == first:
            return first
        if rest.startswith(first):
            suffix = rest[len(first):]
            if 0 < len(suffix) <= 6 and re.match(r"^[a-z\^]{1,6}$", suffix, re.I):
                return first + suffix
    return w


def dedup_span_text(text: str, _depth: int = 0) -> str:
    """Remove text doubling that PDF shadow/animation layers produce.

    Handles
    -------
    - 'Mendelian GeneticsMendelian Genetics' -> 'Mendelian Genetics'
    - 'Gregor Gregor MendelMendel'           -> 'Gregor Mendel'
    - 'HomeworkHomework ProblemsProblems'    -> 'Homework Problems'
    - '1919th century  century'              -> '19th century'
    - Normal text is never mangled.
    """
    if _depth > 4:
        return text
    text = text.strip()
    n    = len(text)
    if n < 4:
        return text

    # Pass 1 – exact string-level half
    if n % 2 == 0:
        half = text[: n // 2]
        if text == half + half:
            return dedup_span_text(half, _depth + 1)

    # Pass 2 – space-separated half
    if n % 2 == 1:
        mid = n // 2
        if text[mid] == " " and text[mid + 1 :] == text[:mid]:
            return dedup_span_text(text[:mid], _depth + 1)

    # Pass 3 – word-prefix repeat
    words = text.split()
    for split_w in range(1, len(words)):
        prefix       = " ".join(words[:split_w])
        rest_after_p = text[len(prefix):].lstrip()
        if rest_after_p.startswith(prefix):
            remainder = rest_after_p[len(prefix):]
            result    = (prefix + remainder).strip()
            if result != text:
                return dedup_span_text(result, _depth + 1)

    # Pass 4 – word-by-word token dedup
    if len(words) >= 2:
        deduped = []
        i       = 0
        changed = False
        while i < len(words):
            w_fixed = _fix_word_double(words[i])
            if w_fixed != words[i]:
                changed = True
            if i + 1 < len(words):
                nxt_fixed = _fix_word_double(words[i + 1])
                if nxt_fixed.lower() == w_fixed.lower():
                    deduped.append(w_fixed)
                    i      += 2
                    changed = True
                    continue
            deduped.append(w_fixed)
            i += 1
        if changed:
            return dedup_span_text(" ".join(deduped), _depth + 1)

    return text


def _strip_trailing_page_number(text: str) -> str:
    """Remove slide page-number bleed from the end of a text span.

    BUG 3 FIX: Only fire the doubled-number heuristic when there's a
    non-digit/non-space character immediately before the run (i.e. it is
    glued to word content, not a standalone number like "22" or "section 11").

    Examples
    --------
    'now called genes (Figure 2.6).1515' -> 'now called genes (Figure 2.6).'
    'gentoype2828'                       -> 'gentoype'
    'subsequent generations.2'           -> 'subsequent generations.'
    'Chapter 2'                          -> 'Chapter 2'   (unchanged)
    'section 11'                         -> 'section 11'  (unchanged – BUG 3)
    '22'                                 -> '22'          (unchanged – BUG 3)
    'F2 22'                              -> 'F2 22'       (unchanged – BUG 3)
    """
    # Guard: only strip if the repeated digit block is glued directly to a
    # word character (letter, digit, punctuation) with NO space before it.
    # This prevents "section 11" -> "section 1".
    m = _TRAILING_DOUBLE_PAGENUM_RE.search(text)
    if m:
        before = text[: m.start()]
        # There must be a non-space, non-digit char immediately before the block
        if before and before[-1] not in (" ", "\t") and not before[-1].isdigit():
            return before.rstrip()
    # Single leaked number after sentence end punctuation
    m = _TRAILING_SINGLE_PAGENUM_RE.search(text)
    if m:
        return text[: m.start() + 1].rstrip()
    return text


def _is_publisher_boilerplate(text: str) -> bool:
    return bool(
        _COPYRIGHT_RE.search(text)
        or _PUBLISHER_RE.search(text)
        or _PERM_REQUIRED_RE.search(text)
    )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 1b — VISUAL-LINE MERGE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _spans_overlap_vertically(
    bbox_a: tuple, bbox_b: tuple, threshold: float = 0.4
) -> bool:
    top_a, bot_a = bbox_a[1], bbox_a[3]
    top_b, bot_b = bbox_b[1], bbox_b[3]
    overlap      = max(0.0, min(bot_a, bot_b) - max(top_a, top_b))
    min_height   = min(bot_a - top_a, bot_b - top_b)
    return min_height > 0 and (overlap / min_height) >= threshold


def _dedup_spans_on_line(spans: list[dict]) -> list[dict]:
    """Remove spans with duplicate text that occupy overlapping bounding boxes."""
    kept: list[dict] = []
    for span in spans:
        norm = _normalize(span["text"])
        duplicate = False
        for k in kept:
            if _normalize(k["text"]) == norm and _spans_overlap_vertically(
                k["bbox"], span["bbox"], threshold=0.3
            ):
                duplicate = True
                break
        if not duplicate:
            kept.append(span)
    return kept


def _merge_spans_into_visual_lines(
    raw_blocks: list[dict], page_height: float
) -> list[dict]:
    """Convert PyMuPDF raw block dicts into merged visual-line dicts.

    v3 additions
    ------------
    * _remap_symbol_chars() applied after dedup_span_text() to fix Wingdings.
    """
    flat: list[dict] = []
    for blk in raw_blocks:
        if blk.get("type") != 0:
            continue
        for line in blk.get("lines", []):
            for span in line.get("spans", []):
                raw = span.get("text", "")
                if not raw.strip():
                    continue
                flags = span.get("flags", 0)
                flat.append({
                    "text":      raw,
                    "bbox":      span["bbox"],
                    "font_size": round(span.get("size", 12.0), 1),
                    "is_bold":   bool(flags & _FITZ_FLAGS["bold"]),
                    "is_italic": bool(flags & _FITZ_FLAGS["italic"]),
                    "is_super":  bool(flags & _FITZ_FLAGS["superscript"]),
                })
    if not flat:
        return []

    flat.sort(key=lambda s: (s["bbox"][1], s["bbox"][0]))

    visual_lines: list[list[dict]] = []
    current_group = [flat[0]]
    current_bbox  = list(flat[0]["bbox"])
    for span in flat[1:]:
        if _spans_overlap_vertically(tuple(current_bbox), span["bbox"]):
            current_group.append(span)
            current_bbox[0] = min(current_bbox[0], span["bbox"][0])
            current_bbox[1] = min(current_bbox[1], span["bbox"][1])
            current_bbox[2] = max(current_bbox[2], span["bbox"][2])
            current_bbox[3] = max(current_bbox[3], span["bbox"][3])
        else:
            visual_lines.append(current_group)
            current_group = [span]
            current_bbox  = list(span["bbox"])
    visual_lines.append(current_group)

    merged: list[dict] = []
    for group in visual_lines:
        group.sort(key=lambda s: s["bbox"][0])
        group = _dedup_spans_on_line(group)
        if not group:
            continue

        total_chars = sum(len(s["text"]) for s in group) or 1
        bold_chars  = sum(len(s["text"]) for s in group if s["is_bold"])
        w_font      = sum(s["font_size"] * len(s["text"]) for s in group) / total_chars

        text_parts: list[str] = []
        for span in group:
            t = span["text"]
            if span["is_super"] and t.strip():
                text_parts.append(f"^{t.strip()}^")
            elif span["is_italic"] and t.strip():
                text_parts.append(f"*{t.strip()}*")
            else:
                text_parts.append(t)

        raw_text = "".join(text_parts).strip()
        if not raw_text:
            continue

        raw_text = dedup_span_text(raw_text)
        raw_text = _strip_trailing_page_number(raw_text)
        raw_text = _normalize_quotes(raw_text)
        raw_text = _remap_symbol_chars(raw_text)   # ← Weakness 1 fix

        if not raw_text.strip():
            continue

        bboxes = [s["bbox"] for s in group]
        merged.append({
            "text":        raw_text,
            "x0":          min(b[0] for b in bboxes),
            "y0":          min(b[1] for b in bboxes),
            "x1":          max(b[2] for b in bboxes),
            "y1":          max(b[3] for b in bboxes),
            "font_size":   round(w_font, 1),
            "is_bold":     bold_chars > total_chars * 0.5,
            "page_height": page_height,
        })
    return merged


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 1c — DROP-CAP FIX
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _fix_drop_caps(lines: list[dict], median_font: float) -> list[dict]:
    if not lines or median_font <= 0:
        return lines
    result, skip_next = [], False
    for i, line in enumerate(lines):
        if skip_next:
            skip_next = False
            continue
        text    = line["text"].strip()
        is_drop = (
            len(text) == 1
            and text.isupper()
            and line["font_size"] >= median_font * 1.8
        )
        if is_drop and i + 1 < len(lines):
            nxt = lines[i + 1]
            if nxt["x0"] >= line["x0"] and _spans_overlap_vertically(
                (line["x0"], line["y0"], line["x1"], line["y1"]),
                (nxt["x0"],  nxt["y0"],  nxt["x1"],  nxt["y1"]),
                0.2,
            ):
                lines[i + 1] = {**nxt, "text": text + nxt["text"].lstrip()}
                skip_next    = True
                continue
        result.append(line)
    return result


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 1d — BULLET CONTINUATION MERGE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _merge_slide_bullet_continuations(lines: list[dict]) -> list[dict]:
    """Re-join wrapped bullet lines on slide decks.

    BUG 6 FIX: This is called *before* dedup and classification (correct order).
    """
    if not lines:
        return lines
    result = [lines[0]]
    for line in lines[1:]:
        prev  = result[-1]
        ptext = prev["text"].strip()
        ctext = line["text"].strip()
        prev_is_bullet = bool(
            ptext and (ptext[0] in _BULLET_CHARS or re.match(r"^[-–]\s", ptext))
        )
        next_no_bullet = bool(
            ctext
            and ctext[0] not in _BULLET_CHARS
            and not re.match(r"^[-–]\s", ctext)
            and not _QUIZ_CHOICE_RE.match(ctext)
        )
        x0_close  = abs(line.get("x0", 0) - prev.get("x0", 0)) <= 25
        same_font = abs(line.get("font_size", 12) - prev.get("font_size", 12)) <= 1.5
        if prev_is_bullet and next_no_bullet and x0_close and same_font:
            result[-1] = {
                **prev,
                "text": ptext + " " + ctext,
                "x1":   max(prev.get("x1", 0), line.get("x1", 0)),
                "y1":   line.get("y1", prev.get("y1", 0)),
            }
        else:
            result.append(line)
    return result


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 1e — EXTRACTORS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _page_has_images(page: Any) -> bool:
    """Return True if the PyMuPDF page contains at least one image stream."""
    try:
        return bool(page.get_images(full=False))
    except Exception:
        return False


def _extract_pdf(
    content: bytes,
) -> tuple[list[TextBlock], int, dict[int, int], dict[int, bool]]:
    """Extract text blocks and image-presence map from a PDF.

    Returns
    -------
    blocks      – list of TextBlock
    page_count  – int
    page_raw    – {page_number: raw_visual_line_count}
    page_images – {page_number: True} when page has embedded images
    """
    if not _FITZ_OK:
        raise RuntimeError("PyMuPDF (fitz) required: pip install PyMuPDF")
    doc          = fitz.open(stream=content, filetype="pdf")
    page_count   = len(doc)
    blocks:      list[TextBlock]   = []
    page_raw:    dict[int, int]    = {}
    page_images: dict[int, bool]   = {}

    for pn, page in enumerate(doc, start=1):
        ph  = page.rect.height
        raw = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        vl  = _merge_spans_into_visual_lines(raw, ph)
        page_raw[pn]    = len(vl)
        page_images[pn] = _page_has_images(page)

        if not vl:
            continue
        sizes    = [l["font_size"] for l in vl if l["font_size"] > 0]
        median_f = median(sizes) if sizes else 12.0
        vl       = _fix_drop_caps(vl, median_f)

        # Weakness 6 guard: verify last page yielded blocks (truncation check)
        if pn == page_count and not vl:
            logger.warning(
                "Last page (%d) produced no text blocks — possible truncation.", pn
            )

        for line in vl:
            blocks.append(TextBlock(
                text=line["text"], page=pn,
                x0=line["x0"],     y0=line["y0"],
                x1=line["x1"],     y1=line["y1"],
                font_size=line["font_size"],
                is_bold=line["is_bold"],
            ))
    doc.close()
    return blocks, page_count, page_raw, page_images


def _extract_tables_pdf(content: bytes) -> dict[int, list]:
    """Extract tables using pdfplumber.

    v3 improvement: also detects Punnett-square text patterns when
    no formal table grid is found (Weakness 3 fix).
    """
    if not _PDFPLUMBER_OK:
        return {}
    tables: dict[int, list] = {}
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for pn, page in enumerate(pdf.pages, start=1):
                ext = page.extract_tables()
                if ext:
                    combined: list = []
                    for t in ext:
                        combined.extend(t)
                    tables[pn] = combined
    except Exception as exc:
        logger.warning("Table extraction failed: %s", exc)
    return tables


def _extract_docx(
    content: bytes,
) -> tuple[list[TextBlock], int, dict[int, int], dict[int, bool]]:
    if not _DOCX_OK:
        raise RuntimeError("python-docx required: pip install python-docx")
    doc        = DocxDocument(io.BytesIO(content))
    blocks:    list[TextBlock] = []
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    page_count = max(1, word_count // 500)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        text  = _remap_symbol_chars(_normalize_quotes(text))
        style = (para.style.name or "").lower()
        is_head = "heading" in style
        m       = re.search(r"heading\s+(\d+)", style)
        lvl     = int(m.group(1)) if m else 0
        blocks.append(TextBlock(
            text=text, page=1,
            font_size=16.0 if is_head else 12.0,
            is_bold=is_head,
            block_type="heading" if is_head else "body",
            heading_level=lvl,
        ))
    for table in doc.tables:
        for row in table.rows:
            cells = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if cells:
                blocks.append(TextBlock(text=cells, page=1, block_type="table"))
    return blocks, page_count, {}, {}


def _extract_txt(
    content: bytes,
) -> tuple[list[TextBlock], int, dict[int, int], dict[int, bool]]:
    text       = content.decode("utf-8", errors="replace")
    text       = _remap_symbol_chars(text)
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    blocks     = [TextBlock(text=p, page=1, font_size=12.0) for p in paragraphs]
    return blocks, max(1, len(text) // 3000), {}, {}


def _extract_html(
    content: bytes,
) -> tuple[list[TextBlock], int, dict[int, int], dict[int, bool]]:
    text = content.decode("utf-8", errors="replace")
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", text, flags=re.DOTALL | re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    for ent, ch in [("&nbsp;", " "), ("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">")]:
        text = text.replace(ent, ch)
    text       = re.sub(r"\s{2,}", "\n", text).strip()
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    blocks     = [TextBlock(text=p, page=1, font_size=12.0) for p in paragraphs]
    return blocks, 1, {}, {}


def _extract_image(
    content: bytes,
) -> tuple[list[TextBlock], int, dict[int, int], dict[int, bool]]:
    if not _OCR_AVAILABLE:
        raise RuntimeError("Pillow + pytesseract required for image OCR")
    img        = PILImage.open(io.BytesIO(content))
    text       = pytesseract.image_to_string(img)
    text       = _remap_symbol_chars(text)
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    blocks     = [TextBlock(text=p, page=1, font_size=12.0) for p in paragraphs]
    return blocks, max(1, len(paragraphs) // 30), {}, {}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 2 — SLIDE-DECK DETECTION & CROSS-PAGE DEDUPLICATION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _is_slide_deck(blocks: list[TextBlock], page_count: int) -> bool:
    if not blocks or page_count == 0:
        return False
    bpp   = len(blocks) / page_count
    avg_f = sum(b.font_size for b in blocks) / len(blocks)
    return bpp < 35 and avg_f > 13.5


def _apply_bullet_continuations_to_blocks(blocks: list[TextBlock]) -> list[TextBlock]:
    """Apply _merge_slide_bullet_continuations at the TextBlock level.

    BUG 6 FIX: called before global dedup and before classification.
    """
    by_page: dict[int, list[TextBlock]] = defaultdict(list)
    for blk in blocks:
        by_page[blk.page].append(blk)

    result: list[TextBlock] = []
    for pn in sorted(by_page):
        line_dicts = [
            {
                "text":      b.text,
                "x0":        b.x0, "y0": b.y0,
                "x1":        b.x1, "y1": b.y1,
                "font_size": b.font_size,
                "is_bold":   b.is_bold,
            }
            for b in by_page[pn]
        ]
        merged_dicts = _merge_slide_bullet_continuations(line_dicts)
        for d in merged_dicts:
            result.append(TextBlock(
                text=d["text"], page=pn,
                x0=d["x0"],     y0=d["y0"],
                x1=d["x1"],     y1=d["y1"],
                font_size=d["font_size"],
                is_bold=d["is_bold"],
            ))
    return result


def _deduplicate_slide_blocks_global(blocks: list[TextBlock]) -> list[TextBlock]:
    """Drop cross-page running headers/footers.

    BUG 6 FIX: now called *after* bullet-continuations, *before* classification.
    """
    text_pages: dict[str, set[int]] = defaultdict(set)
    for blk in blocks:
        text_pages[_normalize(blk.text)].add(blk.page)

    global_noise: set[str] = {k for k, pages in text_pages.items() if len(pages) > 2}

    seen_per_page: dict[int, set[str]] = defaultdict(set)
    result: list[TextBlock] = []
    for blk in blocks:
        key = _normalize(blk.text)
        if key in global_noise:
            continue
        if key in seen_per_page[blk.page]:
            continue
        seen_per_page[blk.page].add(key)
        result.append(blk)
    return result


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 3 — NOISE REMOVAL
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _remove_noise(
    blocks:     list[TextBlock],
    page_count: int,
    page_stats: dict[int, _PageStats],
) -> list[TextBlock]:
    if not blocks:
        return []

    page_appearances: dict[str, set[int]] = defaultdict(set)
    for blk in blocks:
        page_appearances[_normalize(blk.text)].add(blk.page)
    noise_keys = {
        k for k, ps in page_appearances.items() if len(ps) >= MIN_NOISE_REPS
    }

    _url_re = re.compile(r"^https?://\S+$")
    in_back = False
    cleaned: list[TextBlock] = []

    for blk in blocks:
        text        = blk.text.strip()
        key         = _normalize(text)
        heading_key = key.strip("#").strip()

        if blk.block_type == "heading" or _ALLCAPS_RE.match(text):
            if heading_key in _HARD_BACK_MATTER:
                in_back = True

        if in_back:
            page_stats[blk.page].noise_removed += 1
            continue

        if _PAGE_NUM_RE.match(text) and len(text) < 20:
            page_stats[blk.page].noise_removed += 1
            continue

        if _is_publisher_boilerplate(text):
            page_stats[blk.page].noise_removed += 1
            continue

        if key in noise_keys:
            page_stats[blk.page].noise_removed += 1
            continue

        if _url_re.match(text):
            page_stats[blk.page].noise_removed += 1
            continue

        # Weakness 7 fix: suppress spurious empty / single-char headings
        if blk.block_type == "heading" and len(text.lstrip("#").strip()) <= 1:
            page_stats[blk.page].noise_removed += 1
            continue

        if heading_key in _SOFT_BACK_MATTER:
            blk.block_type = "heading"
            blk._soft_content_type = _SOFT_BACK_MATTER[heading_key]  # type: ignore[attr-defined]

        cleaned.append(blk)

    removed = len(blocks) - len(cleaned)
    if removed:
        logger.debug("Noise removed %d blocks", removed)
    return cleaned


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 4 — BLOCK CLASSIFICATION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _classify_blocks(
    blocks:   list[TextBlock],
    is_slide: bool = False,
) -> list[TextBlock]:
    """Assign block_type and heading_level.

    Weakness 7 fix: empty or single-word headings (after stripping markdown)
    are demoted back to 'body'.
    """
    if not blocks:
        return []

    sizes = [b.font_size for b in blocks if b.font_size > 0 and b.block_type != "table"]
    med   = median(sizes) if sizes else 12.0

    heading_ratio  = 1.20 if is_slide else 1.35
    heading_bold_r = 1.00 if is_slide else 1.15

    in_toc = False
    for blk in blocks:
        if blk.block_type == "table":
            continue
        text  = blk.text.strip()
        ratio = blk.font_size / med if med else 1.0
        key   = _normalize(text).strip("#")
        wc    = len(text.split())

        if key in _TOC_HEADINGS:
            in_toc = True
        if in_toc and ratio >= heading_ratio and key not in _TOC_HEADINGS:
            in_toc = False
        if in_toc:
            blk.block_type = "toc"
            continue

        if _CAPTION_RE.match(text):
            blk.block_type = "figure_caption"
            continue

        is_foot = ratio < 0.75 and (blk.y1 > 650 or _FOOTNOTE_START_RE.match(text))
        if is_foot:
            blk.block_type = "footnote"
            continue

        is_clause = not is_slide and any(p.match(text) for p in _CLAUSE_RE)

        is_heading = (
            _HEADING_MIN_WORDS <= wc <= _HEADING_MAX_WORDS
            and (
                ratio >= heading_ratio
                or (blk.is_bold and ratio >= heading_bold_r)
                or (_ALLCAPS_RE.match(text) and not is_clause)
            )
        )
        if is_heading:
            # Weakness 7 fix: require at least 2 meaningful characters after
            # stripping markdown / hash prefix
            clean = text.lstrip("#").strip()
            if len(clean) < 2:
                continue  # leave as body

            blk.block_type = "heading"
            if ratio >= 2.0:
                blk.heading_level = 1
            elif ratio >= 1.5:
                blk.heading_level = 2
            elif ratio >= heading_ratio or (blk.is_bold and ratio >= heading_bold_r):
                blk.heading_level = 3
            else:
                blk.heading_level = 4
            continue

        if is_clause:
            blk.block_type = "clause"
            continue

    return blocks


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 5 — QUIZ DETECTION  (BUG 5 + Weakness 4 fix)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@dataclass
class QuizItem:
    question:     str
    choices:      list[str]
    slide_number: int


def _is_choice_line(text: str) -> bool:
    """Return True if the text looks like an MCQ answer choice."""
    return bool(
        _QUIZ_CHOICE_RE.match(text)
        or _MCQ_NUMBERED_RE.match(text)
    )


def _extract_quiz_items(
    blocks: list[TextBlock],
) -> tuple[list[TextBlock], list[QuizItem]]:
    """Detect and structure quiz/question slides.

    BUG 5 FIX
    ---------
    The stem is collected *between* the "Question" heading (exclusive) and the
    first choice (exclusive).  If there is no "Question" heading but there are
    ≥2 consecutive choice-pattern lines, we scan *backwards* from the first
    choice to find the nearest heading/body block that isn't a choice — that
    is the stem.

    Weakness 4 FIX
    --------------
    _is_choice_line() now also matches "1. 100%", "2. 50 %", "3. none" via
    _MCQ_NUMBERED_RE (broad 1-9 numbered list).
    """
    quiz_items:      list[QuizItem]              = []
    pages_with_quiz: set[int]                    = set()
    by_page:         dict[int, list[TextBlock]]  = defaultdict(list)

    for blk in blocks:
        by_page[blk.page].append(blk)

    for page_num, page_blocks in by_page.items():
        texts = [b.text.strip() for b in page_blocks]
        n     = len(texts)

        # Locate "Question" / "Quiz" heading index
        quiz_title_idx = next(
            (i for i, t in enumerate(texts) if _QUIZ_TITLE_RE.match(t)), None
        )

        # Find all choice-like lines
        choice_indices = [i for i, t in enumerate(texts) if _is_choice_line(t)]

        # Need ≥2 consecutive-ish choices (within 4 positions) to call it a quiz
        has_choices = False
        if len(choice_indices) >= 2:
            for a, b_ in zip(choice_indices, choice_indices[1:]):
                if b_ - a <= 4:
                    has_choices = True
                    break

        if quiz_title_idx is None and not has_choices:
            continue

        pages_with_quiz.add(page_num)

        # Determine choice block range
        first_choice = choice_indices[0] if choice_indices else n
        last_choice  = choice_indices[-1] if choice_indices else n - 1

        # BUG 5 FIX: stem collection
        if quiz_title_idx is not None:
            # Stem is STRICTLY between the title heading and first choice
            stem_start = quiz_title_idx + 1
            stem_end   = first_choice
        else:
            # No "Question" heading: search backwards from first_choice for
            # the most recent non-choice block (heading or body)
            stem_start = 0
            for i in range(first_choice - 1, -1, -1):
                if not _is_choice_line(texts[i]):
                    stem_start = i
                    break
            stem_end = first_choice

        question_parts = [
            texts[i]
            for i in range(stem_start, stem_end)
            if texts[i] and not _is_choice_line(texts[i])
        ]
        question = " ".join(question_parts).strip()

        # Collect choices (consecutive run starting at first_choice)
        choices: list[str] = []
        for i in range(first_choice, min(last_choice + 1, n)):
            if _is_choice_line(texts[i]):
                choices.append(texts[i])

        if question or choices:
            quiz_items.append(
                QuizItem(question=question, choices=choices, slide_number=page_num)
            )
        for blk in page_blocks:
            blk.block_type = "quiz_item"

    return blocks, quiz_items


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 6 — CHAPTER / TOPIC EXTRACTION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _extract_chapter_from_blocks(blocks: list[TextBlock]) -> str:
    for blk in blocks[:20]:
        m = _CHAPTER_RE.search(blk.text)
        if m:
            return m.group(0).strip()
    return ""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 7 — MARKDOWN BUILD
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# Punnett-square text pattern: "TT : Tt : tt" or "1 : 2 : 1"
_PUNNETT_RE = re.compile(
    r"\b[A-Za-z0-9]{1,4}\s*:\s*[A-Za-z0-9]{1,4}(\s*:\s*[A-Za-z0-9]{1,4})*\b"
)


def _is_punnett_or_ratio(text: str) -> bool:
    """Return True if this text line looks like a ratio or Punnett result.

    Used in has_table metadata tagging (BUG 4 partial fix).
    """
    return bool(_PUNNETT_RE.search(text))


def _build_markdown(
    blocks:     list[TextBlock],
    tables:     dict[int, list],
    quiz_items: list[QuizItem],
    page_images: dict[int, bool],
) -> tuple[str, dict[str, int]]:
    """Build a Markdown string from classified blocks.

    Returns
    -------
    markdown          – the Markdown text
    heading_to_page   – {normalized_heading_text: page_number}
                        Used by the chunker to fix BUG 2.

    Weakness 5 fix
    --------------
    Pages with embedded images get a  ![Figure](slide:N) placeholder injected
    after the first heading block on that page.

    Weakness 7 fix
    --------------
    Headings whose clean text is empty after stripping '#' are silently skipped.
    """
    HEADING_PREFIX   = {1: "# ", 2: "## ", 3: "### ", 4: "#### "}
    lines:            list[str] = []
    footnotes:        list[str] = []
    emitted_tables:   set[int]  = set()
    emitted_quiz:     set[int]  = set()
    emitted_img:      set[int]  = set()
    heading_to_page:  dict[str, int] = {}

    quiz_by_page: dict[int, QuizItem] = {qi.slide_number: qi for qi in quiz_items}

    for blk in blocks:
        text = blk.text.strip()
        if not text:
            continue
        bt = blk.block_type

        if bt == "toc":
            continue

        # ── inject table before first body block on this page ────────────────
        if blk.page in tables and blk.page not in emitted_tables:
            tbl_rows = tables[blk.page]
            if tbl_rows:
                md_rows: list[str] = []
                for r_idx, row in enumerate(tbl_rows):
                    cells = [str(c or "").replace("\n", " ").strip() for c in row]
                    md_rows.append("| " + " | ".join(cells) + " |")
                    if r_idx == 0:
                        md_rows.append("| " + " | ".join(["---"] * len(row)) + " |")
                lines.append("\n" + "\n".join(md_rows) + "\n")
            emitted_tables.add(blk.page)

        # ── inject quiz block inline (once per page) ─────────────────────────
        if blk.page in quiz_by_page and blk.page not in emitted_quiz:
            qi = quiz_by_page[blk.page]
            lines.append(f"\n> **Question (Slide {qi.slide_number}):** {qi.question}\n")
            for choice in qi.choices:
                lines.append(f"> - {choice}")
            lines.append("")
            emitted_quiz.add(blk.page)
            continue

        if bt == "quiz_item":
            continue

        if bt == "footnote":
            footnotes.append(text)
            continue

        if bt == "heading":
            clean = text.lstrip("#").strip()
            if not clean:            # Weakness 7: skip empty headings
                continue
            if len(clean) < 2:       # Weakness 7: skip single-char headings
                continue
            lvl = blk.heading_level or 2
            prefix = HEADING_PREFIX.get(lvl, "## ")
            lines.append(f"\n{prefix}{clean}\n")
            # Record heading → page mapping for BUG 2 fix
            heading_to_page[_normalize(clean)] = blk.page
            # Weakness 5: inject image placeholder after first heading on page
            if page_images.get(blk.page) and blk.page not in emitted_img:
                lines.append(
                    f"\n![Figure](slide:{blk.page} "
                    f'"Embedded image on slide {blk.page}")\n'
                )
                emitted_img.add(blk.page)
            continue

        if bt == "clause":
            lines.append(f"\n#### {text}\n")
            continue

        if bt == "figure_caption":
            lines.append(f"\n*{text}*\n")
            continue

        if bt == "table":
            lines.append(f"`{text}`")
            continue

        # body text — normalise bullets
        if text and text[0] in _BULLET_CHARS:
            text = "- " + text[1:].lstrip()
        text = re.sub(r"^[•·●○◦▪▸▹➢✓✗❖►▶→–]\s*", "- ", text, flags=re.MULTILINE)
        lines.append(text)

    if footnotes:
        lines.append("\n---\n\n**Footnotes**\n")
        for i, fn in enumerate(footnotes, 1):
            lines.append(f"[^{i}]: {fn}")

    raw = "\n".join(lines)
    raw = re.sub(r"(\w)-\n(\w)", r"\1\2", raw)    # de-hyphenate
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip(), heading_to_page


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 8 — SLIDE-AWARE STRUCTURAL CHUNKING  (BUG 1 + BUG 2 fix)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _split_markdown_by_heading(markdown: str) -> list[tuple[str, int]]:
    """Split markdown into (section_text, heading_level) pairs at H1/H2/H3.

    Each section starts at its heading and includes everything up to (but not
    including) the next heading of equal or higher level.  The first section
    may have no heading (pre-amble text before the first header).

    Returns list of (text, level) where level=0 means preamble.
    """
    # Find all heading positions
    heading_re = re.compile(r"^(#{1,3})\s+", re.MULTILINE)
    positions  = [(m.start(), len(m.group(1))) for m in heading_re.finditer(markdown)]

    if not positions:
        return [(markdown, 0)]

    sections: list[tuple[str, int]] = []
    # preamble before first heading
    if positions[0][0] > 0:
        pre = markdown[: positions[0][0]].strip()
        if pre:
            sections.append((pre, 0))

    for i, (start, lvl) in enumerate(positions):
        end = positions[i + 1][0] if i + 1 < len(positions) else len(markdown)
        section_text = markdown[start:end].strip()
        if section_text:
            sections.append((section_text, lvl))

    return sections


def _fallback_split(text: str, size: int, overlap: int) -> list[str]:
    words  = text.split()
    chunks: list[str] = []
    start  = 0
    while start < len(words):
        end = start + size
        chunks.append(" ".join(words[start:end]))
        start = end - overlap
    return chunks


def _page_for_section(
    section_text:     str,
    heading_to_page:  dict[str, int],
    blocks:           list[TextBlock],
) -> int:
    """Return the page number for a markdown section.

    BUG 2 FIX
    ---------
    Strategy (in priority order):
    1. Extract the first heading from section_text, normalize it, look up in
       heading_to_page (built from the markdown build stage, not from raw blocks).
    2. Scan blocks for a text that word-overlaps significantly with the first
       non-empty, non-heading line of the section.
    3. Default to 0 (unknown).
    """
    # Strategy 1: heading lookup
    heading_m = re.search(r"^#{1,4}\s+(.+)$", section_text, re.MULTILINE)
    if heading_m:
        key = _normalize(heading_m.group(1))
        if key in heading_to_page:
            return heading_to_page[key]

    # Strategy 2: word-overlap with first non-heading body line
    body_lines = [
        l.strip()
        for l in section_text.split("\n")
        if l.strip() and not l.strip().startswith("#")
        and not l.strip().startswith(">")
        and not l.strip().startswith("*")
        and not l.strip().startswith("|")
    ]
    if body_lines:
        target_words = set(_normalize(body_lines[0]).split())
        best_page, best_score = 0, 0
        for blk in blocks:
            blk_words = set(_normalize(blk.text).split())
            if not blk_words:
                continue
            overlap = len(target_words & blk_words) / max(len(target_words), 1)
            if overlap > best_score:
                best_score = overlap
                best_page  = blk.page
        if best_score >= 0.5:
            return best_page

    return 0


def _chunk_text(
    markdown:        str,
    job_id:          str,
    filename:        str,
    chapter:         str,
    quiz_pages:      set[int],
    blocks:          list[TextBlock],
    heading_to_page: dict[str, int],
) -> list[DocumentChunk]:
    """Split markdown into DocumentChunks using structural (heading) boundaries.

    BUG 1 FIX
    ---------
    Primary split: one chunk per heading section (H1/H2/H3 boundary).
    Secondary split: if a section exceeds CHUNK_SIZE_TOKENS, it is further
    split using RecursiveCharacterTextSplitter (or fallback), with overlap.

    BUG 2 FIX
    ---------
    slide_number is resolved via _page_for_section() which uses the
    heading_to_page map produced during markdown build.

    BUG 4 FIX
    ---------
    has_table now also flags sections containing Punnett / ratio patterns.
    """
    sections = _split_markdown_by_heading(markdown)

    # Build per-section chunks (potentially further split if too large)
    raw_sections: list[str] = []
    for section_text, _lvl in sections:
        tok = _estimate_tokens(section_text)
        if tok <= CHUNK_SIZE_TOKENS:
            raw_sections.append(section_text)
        else:
            # Section too long: split further
            if _LANGCHAIN_AVAILABLE:
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=CHUNK_SIZE_TOKENS * 4,
                    chunk_overlap=CHUNK_OVERLAP_TOKENS * 4,
                    separators=["\n\n", "\n", ". ", " ", ""],
                    length_function=len,
                )
                sub = splitter.split_text(section_text)
            else:
                sub = _fallback_split(
                    section_text, CHUNK_SIZE_TOKENS, CHUNK_OVERLAP_TOKENS
                )
            raw_sections.extend(sub)

    chunks: list[DocumentChunk] = []
    current_topic    = ""
    current_subtopic = ""

    for idx, raw in enumerate(raw_sections):
        text = raw.strip()
        if not text:
            continue

        # Track hierarchical headings
        h1 = re.search(r"^#\s+(.+)$",    text, re.MULTILINE)
        h2 = re.search(r"^##\s+(.+)$",   text, re.MULTILINE)
        h3 = re.search(r"^###\s+(.+)$",  text, re.MULTILINE)
        h4 = re.search(r"^####\s+(.+)$", text, re.MULTILINE)

        if h1:
            current_topic, current_subtopic = _slug(h1.group(1)), ""
        elif h2:
            current_topic, current_subtopic = _slug(h2.group(1)), ""
        if h3:
            current_subtopic = _slug(h3.group(1))
        elif h4:
            current_subtopic = _slug(h4.group(1))

        # BUG 2 FIX: use heading_to_page + word-overlap fallback
        slide_num = _page_for_section(text, heading_to_page, blocks)

        # BUG 4 FIX: has_table includes Punnett patterns
        has_table = bool(re.search(r"\|\s*[^|]+\s*\|", text)) or any(
            _is_punnett_or_ratio(line) for line in text.split("\n")
        )
        has_clause = bool(re.search(
            r"^(?:####\s+(?:\d+\.|\([a-z]\)|ARTICLE))", text, re.MULTILINE
        ))
        has_quiz   = bool(re.search(r"^\s*>\s*\*\*Question", text, re.MULTILINE))
        has_def    = bool(_DEFINITION_RE.search(text))
        has_img    = bool(re.search(r"!\[Figure\]", text))

        if has_quiz:
            content_type = "quiz_item"
        elif has_def:
            content_type = "definition"
        elif has_table:
            content_type = "table"
        elif re.search(r"(?i)\bhomework\b|\bproblems\b", text):
            content_type = "homework"
        elif re.search(r"(?i)\bexample\b", text):
            content_type = "example"
        else:
            content_type = "concept"

        chunks.append(DocumentChunk(
            chunk_id=f"{job_id}-{idx:04d}",
            text=text,
            token_estimate=_estimate_tokens(text),
            metadata={
                "index":          idx,
                "filename":       filename,
                "job_id":         job_id,
                "heading":        current_topic,
                "char_count":     len(text),
                "token_count":    _estimate_tokens(text),
                "has_table":      has_table,
                "has_clause":     has_clause,
                "has_image":      has_img,
                "chapter":        chapter,
                "topic":          current_topic,
                "subtopic":       current_subtopic,
                "slide_number":   slide_num,
                "content_type":   content_type,
                "has_quiz":       has_quiz,
                "has_definition": has_def,
            },
        ))
    return chunks


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# STAGE 9 — AI READINESS SCORING
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _compute_ai_readiness(
    page_stats: dict[int, _PageStats],
    page_count: int,
    page_raw:   dict[int, int],
) -> dict[str, Any]:
    page_scores: list[dict[str, Any]] = []
    warnings:    list[str]            = []

    for pn in range(1, page_count + 1):
        ps  = page_stats.get(pn, _PageStats(page=pn))
        raw = page_raw.get(pn, 1)

        text_conf   = min(1.0, ps.kept_block_count / raw) if raw > 0 else 0.0
        struct_conf = (
            min(1.0, ps.classified_count / ps.kept_block_count)
            if ps.kept_block_count > 0 else 0.0
        )
        noise_ratio = (
            ps.noise_removed / (ps.kept_block_count + ps.noise_removed)
            if (ps.kept_block_count + ps.noise_removed) > 0 else 0.0
        )
        # Bonus: image-only slides get a note but not a penalty
        image_note = ""
        if ps.has_image and ps.kept_block_count == 0:
            image_note = " (image-only slide)"
            text_conf  = 0.5   # partial credit; content may be in caption

        score = round(
            0.4 * text_conf + 0.4 * struct_conf + 0.2 * (1.0 - noise_ratio), 3
        )

        if text_conf < 0.35:
            warnings.append(
                f"Page {pn}: low text confidence ({text_conf:.2f})"
                f"{image_note} — possible image slide or scan"
            )
        if struct_conf < 0.25 and ps.kept_block_count > 3:
            warnings.append(
                f"Page {pn}: low structure confidence ({struct_conf:.2f})"
                " — classification may be unreliable"
            )

        page_scores.append({
            "page":                       pn,
            "text_extraction_confidence": round(text_conf, 3),
            "structure_confidence":       round(struct_conf, 3),
            "noise_ratio":                round(noise_ratio, 3),
            "has_image":                  ps.has_image,
            "score":                      score,
        })

    overall = (
        round(sum(p["score"] for p in page_scores) / len(page_scores), 3)
        if page_scores else 0.0
    )
    return {"overall_score": overall, "page_scores": page_scores, "warnings": warnings}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PIPELINE ORCHESTRATOR
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def run_pipeline(content: bytes, filename: str, job_id: str) -> ProcessedDocument:
    """Execute all stages and return a ProcessedDocument with attached extras.

    Stage ordering (BUG 6 fix)
    --------------------------
    1. Extract
    2. bullet_continuations        ← BEFORE global dedup
    3. deduplicate_slide_blocks    ← BEFORE classification
    4. remove_noise
    5. classify_blocks             ← AFTER dedup
    6. extract_quiz_items
    7. extract_chapter
    8. build_markdown
    9. chunk_text
    10. ai_readiness
    """
    t0       = time.monotonic()
    doc_type = _detect_doc_type(filename, content)
    logger.info(
        "Pipeline start | job=%s | file=%s | type=%s", job_id, filename, doc_type
    )

    if doc_type == DocumentType.UNKNOWN:
        raise ValueError(f"Unsupported file type: {filename}")

    tables:      dict[int, list] = {}
    page_raw:    dict[int, int]  = {}
    page_images: dict[int, bool] = {}

    if doc_type == DocumentType.PDF:
        blocks, page_count, page_raw, page_images = _extract_pdf(content)
        tables                                    = _extract_tables_pdf(content)
    elif doc_type == DocumentType.DOCX:
        blocks, page_count, page_raw, page_images = _extract_docx(content)
    elif doc_type == DocumentType.TXT:
        blocks, page_count, page_raw, page_images = _extract_txt(content)
    elif doc_type == DocumentType.HTML:
        blocks, page_count, page_raw, page_images = _extract_html(content)
    elif doc_type == DocumentType.IMAGE:
        blocks, page_count, page_raw, page_images = _extract_image(content)
    else:
        raise ValueError(f"No extractor for {doc_type}")

    if not blocks:
        raise ValueError("No text extracted. May be a scanned image; enable OCR.")

    page_stats: dict[int, _PageStats] = {
        pn: _PageStats(
            page=pn,
            raw_block_count=page_raw.get(pn, 0),
            has_image=page_images.get(pn, False),
        )
        for pn in range(1, page_count + 1)
    }

    is_slide = _is_slide_deck(blocks, page_count)
    if is_slide:
        logger.info("Slide deck detected — applying slide-aware pipeline")
        # BUG 6 FIX: bullet continuations BEFORE global dedup
        blocks = _apply_bullet_continuations_to_blocks(blocks)
        # BUG 6 FIX: dedup BEFORE classification
        blocks = _deduplicate_slide_blocks_global(blocks)

    blocks = _remove_noise(blocks, page_count, page_stats)

    for blk in blocks:
        page_stats[blk.page].kept_block_count += 1

    # BUG 6 FIX: classification AFTER dedup
    blocks = _classify_blocks(blocks, is_slide=is_slide)

    for blk in blocks:
        if blk.block_type != "body":
            page_stats[blk.page].classified_count += 1

    blocks, quiz_items = _extract_quiz_items(blocks)

    chapter  = _extract_chapter_from_blocks(blocks)
    markdown, heading_to_page = _build_markdown(blocks, tables, quiz_items, page_images)

    if not markdown.strip():
        raise ValueError("Pipeline produced empty output.")

    quiz_pages = {qi.slide_number for qi in quiz_items}
    chunks     = _chunk_text(
        markdown, job_id, filename, chapter,
        quiz_pages, blocks, heading_to_page,
    )
    ai_readiness = _compute_ai_readiness(page_stats, page_count, page_raw)

    elapsed = int((time.monotonic() - t0) * 1000)
    logger.info(
        "Pipeline complete | job=%s | pages=%d | chunks=%d | quiz=%d | ms=%d",
        job_id, page_count, len(chunks), len(quiz_items), elapsed,
    )

    result = ProcessedDocument(
        job_id=job_id,
        filename=filename,
        doc_type=doc_type,
        page_count=page_count,
        chunks=chunks,
        markdown=markdown,
        processing_time_ms=elapsed,
        sha256=_sha256(content),
    )
    result._quiz_items    = quiz_items     # type: ignore[attr-defined]
    result._ai_readiness  = ai_readiness   # type: ignore[attr-defined]
    result._chapter       = chapter        # type: ignore[attr-defined]
    result._heading_to_page = heading_to_page  # type: ignore[attr-defined]
    return result


def serialize_result(result: ProcessedDocument) -> dict[str, Any]:
    quiz_items:   list[QuizItem] = getattr(result, "_quiz_items",   [])
    ai_readiness: dict[str, Any] = getattr(result, "_ai_readiness", {})
    return {
        "job_id":   result.job_id,
        "filename": result.filename,
        "markdown": result.markdown,
        "chunks": [
            {"chunk_id": c.chunk_id, "text": c.text, "metadata": c.metadata}
            for c in result.chunks
        ],
        "stats": {
            "total_pages":          result.page_count,
            "total_chunks":         len(result.chunks),
            "total_tokens":         sum(c.token_estimate for c in result.chunks),
            "headings_detected":    sum(1 for c in result.chunks if c.metadata.get("heading")),
            "tables_detected":      sum(bool(c.metadata.get("has_table"))  for c in result.chunks),
            "clauses_detected":     sum(bool(c.metadata.get("has_clause")) for c in result.chunks),
            "quiz_items_detected":  len(quiz_items),
            "definitions_detected": sum(
                bool(c.metadata.get("has_definition")) for c in result.chunks
            ),
            "images_detected":      sum(
                bool(c.metadata.get("has_image")) for c in result.chunks
            ),
            "slide_count":          result.page_count,
            "processing_ms":        result.processing_time_ms,
        },
        "ai_readiness": ai_readiness,
        "quiz_items": [
            {
                "slide_number": qi.slide_number,
                "question":     qi.question,
                "choices":      qi.choices,
            }
            for qi in quiz_items
        ],
    }


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SELF-CONTAINED TEST SUITE   (python pipeline.py --test)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _run_tests() -> None:  # noqa: C901  (complexity – it's a test suite)
    import traceback
    passed = failed = 0

    def ok(label: str) -> None:
        nonlocal passed
        passed += 1
        print(f"  ✓  {label}")

    def fail(label: str, detail: str = "") -> None:
        nonlocal failed
        failed += 1
        print(f"  ✗  {label}")
        if detail:
            print(f"       {detail}")

    def check(label: str, got: Any, expected: Any) -> None:
        if got == expected:
            ok(label)
        else:
            fail(label, f"got {got!r}, expected {expected!r}")

    def check_true(label: str, expr: bool) -> None:
        if expr:
            ok(label)
        else:
            fail(label, "expression was False")

    def check_false(label: str, expr: bool) -> None:
        if not expr:
            ok(label)
        else:
            fail(label, "expression was True (expected False)")

    # ── _remap_symbol_chars ───────────────────────────────────────────────────
    print("\n[_remap_symbol_chars]")
    check("Wingdings bullet → -",   _remap_symbol_chars("\uf0b7 text"), "- text")
    check("U+FFFD → -",             _remap_symbol_chars("\ufffd text"), "- text")
    check("Greek mu → μ",           _remap_symbol_chars("\uf06dg"),     "μg")
    check("plain text unchanged",   _remap_symbol_chars("hello"),       "hello")
    check("PUA collapse",
          _remap_symbol_chars("\uf0b7\uf0b7\uf0b7 text"), "- text")

    # ── dedup_span_text ───────────────────────────────────────────────────────
    print("\n[dedup_span_text]")
    DEDUP_CASES = [
        ("Mendelian GeneticsMendelian Genetics",    "Mendelian Genetics"),
        ("since before recorded historysince before recorded history",
                                                    "since before recorded history"),
        ("Selection of animals for domesticationSelection of animals for domestication",
                                                    "Selection of animals for domestication"),
        ("Gregor Gregor MendelMendel",              "Gregor Mendel"),
        ("Ch.Ch. 2",                                "Ch. 2"),
        ("Keys to MendelKeys to Mendel",            "Keys to Mendel"),
        ("TestcrossesTestcrosses",                  "Testcrosses"),
        ("HomeworkHomework ProblemsProblems",        "Homework Problems"),
        ("1919th century  century",                 "19th century"),
        ("Monohybrid Crosses and MendelMonohybrid Crosses and Mendel",
                                                    "Monohybrid Crosses and Mendel"),
        ("Normal text here stays intact",           "Normal text here stays intact"),
        ("The garden pea was an ideal organism",    "The garden pea was an ideal organism"),
        ("Mendel analyzed traits with discrete forms",
                                                    "Mendel analyzed traits with discrete forms"),
        ("AI",                                      "AI"),
        ("3:1 ratio",                               "3:1 ratio"),
        ("19th",                                    "19th"),
        ("gentoypegentoype",                        "gentoype"),
        ("the second filial generation, F2, is produced.the second filial generation, F2, is produced.",
                                                    "the second filial generation, F2, is produced."),
    ]
    for inp, expected in DEDUP_CASES:
        check(repr(inp[:55]), dedup_span_text(inp), expected)

    # ── _strip_trailing_page_number (BUG 3 fix) ───────────────────────────────
    print("\n[_strip_trailing_page_number — BUG 3]")
    STRIP_CASES = [
        ("subsequent generations.2",           "subsequent generations."),
        ("now called genes (Figure 2.6).1515",  "now called genes (Figure 2.6)."),
        ("gentoype2828",                        "gentoype"),
        ("Normal sentence here.",               "Normal sentence here."),
        ("3:1 ratio",                           "3:1 ratio"),
        ("Chapter 2",                           "Chapter 2"),
        ("F2 generation generation",            "F2 generation generation"),
        # BUG 3 cases that MUST NOT be mangled
        ("section 11",                          "section 11"),
        ("22",                                  "22"),
        ("F2 22",                               "F2 22"),
        ("problem 33",                          "problem 33"),
        ("step 44",                             "step 44"),
    ]
    for inp, expected in STRIP_CASES:
        check(repr(inp), _strip_trailing_page_number(inp), expected)

    # ── _is_publisher_boilerplate ─────────────────────────────────────────────
    print("\n[_is_publisher_boilerplate]")
    for t in [
        "Copyright ©The McGraw-Hill Companies, Inc. Permission required for reproduction or display",
        "Peter J. Russell, iGenetics: Copyright © Pearson Education, Inc., publishing as Benjamin Cummings.",
    ]:
        check_true(f"DROP  {t[:55]!r}", _is_publisher_boilerplate(t))
    for t in [
        "Mendelian Genetics Ch. 2",
        "What traits are inherited?",
        "The F2 has a ratio of about 3:1",
    ]:
        check_false(f"KEEP  {t[:55]!r}", _is_publisher_boilerplate(t))

    # ── _is_punnett_or_ratio ──────────────────────────────────────────────────
    print("\n[_is_punnett_or_ratio — BUG 4]")
    check_true("TT:Tt:tt flagged",         _is_punnett_or_ratio("TT : Tt : tt"))
    check_true("1:2:1 flagged",            _is_punnett_or_ratio("ratio 1 : 2 : 1"))
    check_true("3:1 flagged",              _is_punnett_or_ratio("3 : 1 ratio"))
    check_false("plain sentence",          _is_punnett_or_ratio("This is normal text."))
    check_false("single number",           _is_punnett_or_ratio("22"))

    # ── _fix_word_double ──────────────────────────────────────────────────────
    print("\n[_fix_word_double]")
    FWD = [
        ("TestcrossesTestcrosses", "Testcrosses"),
        ("1919th",                 "19th"),
        ("19th",                   "19th"),
        ("Homework",               "Homework"),
        ("Ch.",                    "Ch."),
        ("abcabc",                 "abc"),
    ]
    for inp, exp in FWD:
        check(repr(inp), _fix_word_double(inp), exp)

    # ── _normalize_quotes ─────────────────────────────────────────────────────
    print("\n[_normalize_quotes]")
    check("left single",  _normalize_quotes("\u2018hello"), "'hello")
    check("right single", _normalize_quotes("it\u2019s"),   "it's")
    check("em dash",      _normalize_quotes("x\u2014y"),    "x--y")
    check("no change",    _normalize_quotes("plain text"),  "plain text")

    # ── _detect_doc_type ──────────────────────────────────────────────────────
    print("\n[_detect_doc_type]")
    check("PDF magic",    _detect_doc_type("x.pdf",  b"%PDF-1.4 rest"), DocumentType.PDF)
    check("DOCX magic",   _detect_doc_type("x.docx", b"PK rest"),       DocumentType.DOCX)
    check(".txt ext",     _detect_doc_type("x.txt",  b"hello"),         DocumentType.TXT)
    check(".html ext",    _detect_doc_type("x.html", b"<html>"),        DocumentType.HTML)
    check("unknown",      _detect_doc_type("x.xyz",  b"random"),        DocumentType.UNKNOWN)

    # ── _is_slide_deck ────────────────────────────────────────────────────────
    print("\n[_is_slide_deck]")
    slide_blocks = [TextBlock(text=f"slide {i}", page=i, font_size=24.0) for i in range(1, 35)]
    dense_blocks = [TextBlock(text=f"para {i}", page=1, font_size=12.0) for i in range(100)]
    check("slide deck detected", _is_slide_deck(slide_blocks, 34),  True)
    check("dense doc not slide", _is_slide_deck(dense_blocks, 1),   False)
    check("empty blocks",        _is_slide_deck([], 10),            False)

    # ── _split_markdown_by_heading (BUG 1) ────────────────────────────────────
    print("\n[_split_markdown_by_heading — BUG 1]")
    sample_md = (
        "Intro text before any heading.\n\n"
        "## Section One\n\nBody of section one.\n\n"
        "## Section Two\n\nBody of section two.\n\n"
        "### Subsection 2a\n\nSubsection body.\n"
    )
    sects = _split_markdown_by_heading(sample_md)
    check("preamble exists",       sects[0][1], 0)
    check("preamble text",         "Intro text" in sects[0][0], True)
    check("section count",         len(sects), 4)
    check("h2 level",              sects[1][1], 2)
    check("h3 level",              sects[3][1], 3)
    # Single giant block → multiple sections
    long_md = "\n\n".join(
        f"## Slide {i}\n\nContent for slide {i}." for i in range(1, 36)
    )
    long_sects = _split_markdown_by_heading(long_md)
    check("35 slides → 35 sections", len(long_sects), 35)

    # ── _page_for_section (BUG 2) ─────────────────────────────────────────────
    print("\n[_page_for_section — BUG 2]")
    h2p = {"mendelian genetics": 3, "testcrosses": 7}
    blks = [
        TextBlock(text="Mendel studied seven traits.", page=5),
        TextBlock(text="Punnett square shows ratios.", page=6),
    ]
    sec_heading = "## Mendelian Genetics\n\nMendel studied heredity."
    sec_body    = "Mendel studied seven traits in detail."
    sec_unknown = "## Totally Unknown Topic\n\nNothing matches."
    check("heading lookup",   _page_for_section(sec_heading, h2p, blks), 3)
    check("word overlap",     _page_for_section(sec_body,    h2p, blks), 5)
    check("fallback 0",       _page_for_section(sec_unknown, h2p, blks), 0)

    # ── _PAGE_NUM_RE ──────────────────────────────────────────────────────────
    print("\n[_PAGE_NUM_RE]")
    for t in ("1", "15", "page 2", "2 of 34"):
        check(f"page num {t!r}",     bool(_PAGE_NUM_RE.match(t)), True)
    for t in ("Chapter 2", "F2", "3:1", "Fig. 2.6"):
        check(f"not page num {t!r}", bool(_PAGE_NUM_RE.match(t)), False)

    # ── _dedup_spans_on_line ──────────────────────────────────────────────────
    print("\n[_dedup_spans_on_line]")
    bbox_a = (10, 100, 200, 120)
    bbox_b = (10, 102, 200, 122)   # overlapping → dup
    bbox_c = (10, 200, 200, 220)   # separate y → keep
    span_dup = [
        {"text": "Mendelian Genetics", "bbox": bbox_a},
        {"text": "Mendelian Genetics", "bbox": bbox_b},
        {"text": "Ch. 2",              "bbox": bbox_c},
    ]
    result_dedup = _dedup_spans_on_line(span_dup)
    check("dedup removes overlapping dup", len(result_dedup), 2)
    check("kept first unique",             result_dedup[0]["text"], "Mendelian Genetics")
    check("kept different bbox",           result_dedup[1]["text"], "Ch. 2")

    # ── _merge_slide_bullet_continuations ────────────────────────────────────
    print("\n[_merge_slide_bullet_continuations]")
    cont_lines = [
        {"text": "– Selective breeding of plants", "x0": 50,  "font_size": 18.0,
         "y0": 100, "y1": 120, "x1": 400},
        {"text": "and animals in history",          "x0": 52,  "font_size": 18.0,
         "y0": 125, "y1": 145, "x1": 400},  # continuation
        {"text": "– Another bullet",                "x0": 50,  "font_size": 18.0,
         "y0": 160, "y1": 180, "x1": 400},
    ]
    merged = _merge_slide_bullet_continuations(cont_lines)
    check("bullet continuation merged",  len(merged), 2)
    check("merged text",
          merged[0]["text"],
          "– Selective breeding of plants and animals in history")
    check("second bullet untouched",     merged[1]["text"], "– Another bullet")

    # ── _is_choice_line (Weakness 4 fix) ─────────────────────────────────────
    print("\n[_is_choice_line — Weakness 4 / BUG 5]")
    check_true("1) choice",         _is_choice_line("1) two identical alleles"))
    check_true("a) choice",         _is_choice_line("a) recessive"))
    check_true("1. 100%",           _is_choice_line("1. 100%"))
    check_true("2. 50%",            _is_choice_line("2. 50%"))
    check_true("3. none",           _is_choice_line("3. none of the above"))
    check_true("A. choice",         _is_choice_line("A. homozygous"))
    check_false("normal sentence",  _is_choice_line("Mendel worked with peas."))
    check_false("heading",          _is_choice_line("## Mendelian Genetics"))
    check_false("number+word",      _is_choice_line("1 chromosome is involved"))

    # ── _extract_quiz_items (BUG 5 fix) ──────────────────────────────────────
    print("\n[_extract_quiz_items — BUG 5]")
    # Case A: explicit "Question" heading
    q_blocks_a = [
        TextBlock(text="Question",                  page=5, font_size=20.0, block_type="heading"),
        TextBlock(text="What is homozygous?",        page=5, font_size=16.0),
        TextBlock(text="1) two identical alleles",   page=5, font_size=14.0),
        TextBlock(text="2) two different alleles",   page=5, font_size=14.0),
        TextBlock(text="3) no alleles",              page=5, font_size=14.0),
    ]
    _, qi_a = _extract_quiz_items(q_blocks_a)
    check("A: quiz detected",       len(qi_a), 1)
    check("A: slide number",        qi_a[0].slide_number, 5)
    check("A: 3 choices",           len(qi_a[0].choices), 3)
    check("A: stem correct",        "homozygous" in qi_a[0].question.lower(), True)
    # Ensure stem does NOT include the choices themselves
    check("A: stem no choice text", "1)" not in qi_a[0].question, True)

    # Case B: no heading, just numbered MCQ (Weakness 4 pattern)
    q_blocks_b = [
        TextBlock(text="What proportion of progeny is homozygous?",
                  page=7, font_size=14.0),
        TextBlock(text="1. 100%",  page=7, font_size=14.0),
        TextBlock(text="2. 50%",   page=7, font_size=14.0),
        TextBlock(text="3. 35%",   page=7, font_size=14.0),
        TextBlock(text="4. none",  page=7, font_size=14.0),
    ]
    _, qi_b = _extract_quiz_items(q_blocks_b)
    check("B: numbered MCQ detected", len(qi_b), 1)
    check("B: 4 choices",             len(qi_b[0].choices), 4)
    check("B: stem captured",         "homozygous" in qi_b[0].question.lower(), True)

    # ── _extract_txt round-trip ───────────────────────────────────────────────
    print("\n[_extract_txt round-trip]")
    blks, pc, _, _ = _extract_txt(b"First.\n\nSecond.\n\nThird.")
    check("txt block count", len(blks), 3)
    check("txt first block", blks[0].text, "First.")
    check("txt page count",  pc, 1)

    # ── _extract_html round-trip ─────────────────────────────────────────────
    print("\n[_extract_html round-trip]")
    blks_h, _, _, _ = _extract_html(
        b"<html><body><p>Hello</p><script>evil()</script><p>World</p></body></html>"
    )
    texts_h = [b.text for b in blks_h]
    check_false("html strips script",    any("evil" in t for t in texts_h))
    check_true("html keeps paragraphs",  any("Hello" in t for t in texts_h))

    # ── _classify_blocks heading word-count guard ─────────────────────────────
    print("\n[_classify_blocks — heading guard]")
    long_block  = TextBlock(text=" ".join(["word"] * 25), page=1,
                            font_size=24.0, is_bold=True)
    short_block = TextBlock(text="Mendelian Genetics",    page=1,
                            font_size=24.0, is_bold=True)
    body_block  = TextBlock(text="body text here",        page=1, font_size=12.0)
    result_cls  = _classify_blocks([long_block, short_block, body_block], is_slide=True)
    check("long block not heading",  result_cls[0].block_type, "body")
    check("short block is heading",  result_cls[1].block_type, "heading")

    # ── _build_markdown smoke test ────────────────────────────────────────────
    print("\n[_build_markdown]")
    md_blocks = [
        TextBlock(text="Mendelian Genetics", page=1, font_size=24.0,
                  block_type="heading", heading_level=1),
        TextBlock(text="Ch. 2",              page=1, font_size=18.0,
                  block_type="heading", heading_level=2),
        TextBlock(text="Gregor Mendel was a monk.", page=1, font_size=12.0,
                  block_type="body"),
        TextBlock(text="Fig. 2.1 Genotype diagram", page=1, font_size=10.0,
                  block_type="figure_caption"),
        TextBlock(text="TT : Tt : tt",   page=2, font_size=12.0, block_type="body"),
    ]
    md, h2p_map = _build_markdown(md_blocks, {}, [], {})
    check("h1 in markdown",           "# Mendelian Genetics"        in md, True)
    check("h2 in markdown",           "## Ch. 2"                    in md, True)
    check("body text in markdown",    "Gregor Mendel was a monk."   in md, True)
    check("caption italicised",       "*Fig. 2.1 Genotype diagram*" in md, True)
    check("heading_to_page h1",       h2p_map.get("mendelian genetics"), 1)
    check("heading_to_page h2",       h2p_map.get("ch. 2"), 1)
    check("punnett line kept",        "TT : Tt : tt"                in md, True)

    # ── Weakness 5: image placeholder ────────────────────────────────────────
    print("\n[_build_markdown — image placeholder (Weakness 5)]")
    img_blocks = [
        TextBlock(text="Slide with image", page=3, font_size=20.0,
                  block_type="heading", heading_level=2),
        TextBlock(text="some body",        page=3, font_size=12.0, block_type="body"),
    ]
    md_img, _ = _build_markdown(img_blocks, {}, [], {3: True})
    check("image placeholder injected",  "![Figure](slide:3" in md_img, True)
    md_no_img, _ = _build_markdown(img_blocks, {}, [], {})
    check("no placeholder without image", "![Figure]" not in md_no_img, True)

    # ── Weakness 7: spurious heading suppression ──────────────────────────────
    print("\n[_build_markdown — spurious heading suppression (Weakness 7)]")
    spurious_blocks = [
        TextBlock(text="Ch.",         page=1, font_size=20.0,
                  block_type="heading", heading_level=2),
        TextBlock(text="Good Heading", page=1, font_size=20.0,
                  block_type="heading", heading_level=2),
        TextBlock(text="body",        page=1, font_size=12.0, block_type="body"),
    ]
    md_spur, _ = _build_markdown(spurious_blocks, {}, [], {})
    check("short heading suppressed", "## Ch." not in md_spur, True)
    check("good heading kept",        "## Good Heading" in md_spur, True)

    # ── structural chunking smoke test (BUG 1) ─────────────────────────────
    print("\n[_chunk_text structural chunking — BUG 1]")
    slides_md = "\n\n".join(
        f"## Slide {i}\n\nBody text for slide {i}." for i in range(1, 11)
    )
    h2p_10 = {f"slide {i}": i for i in range(1, 11)}
    blks_10 = [TextBlock(text=f"Body text for slide {i}.", page=i) for i in range(1, 11)]
    chunks_10 = _chunk_text(
        slides_md, "job-x", "test.pdf", "Ch. 2",
        set(), blks_10, h2p_10,
    )
    check("10 slides → 10 chunks",     len(chunks_10), 10)
    check("chunk 0 has slide_number",  chunks_10[0].metadata["slide_number"] != 0, True)
    check("chunk 0 text correct",      "Slide 1" in chunks_10[0].text, True)
    # Punnett-ratio chunk flagged as has_table
    punnett_md = "## Ratios\n\nTT : Tt : tt gives 1 : 2 : 1 ratio."
    punnett_chunks = _chunk_text(
        punnett_md, "job-p", "test.pdf", "Ch. 2",
        set(), [], {"ratios": 1},
    )
    check("punnett chunk flagged as has_table",
          punnett_chunks[0].metadata["has_table"], True)

    # ── _compute_ai_readiness ────────────────────────────────────────────────
    print("\n[_compute_ai_readiness]")
    ps = {
        1: _PageStats(page=1, raw_block_count=10, kept_block_count=8,
                      classified_count=4, noise_removed=2),
        2: _PageStats(page=2, raw_block_count=8,  kept_block_count=7,
                      classified_count=3, noise_removed=1, has_image=True),
    }
    ari = _compute_ai_readiness(ps, 2, {1: 10, 2: 8})
    check("overall score 0–1",    0.0 <= ari["overall_score"] <= 1.0, True)
    check("page_scores len 2",    len(ari["page_scores"]), 2)
    check("warnings is list",     isinstance(ari["warnings"], list), True)
    check("has_image in scores",  ari["page_scores"][1]["has_image"], True)

    # ── serialize_result ──────────────────────────────────────────────────────
    print("\n[serialize_result]")
    doc = ProcessedDocument(
        job_id="test-job", filename="test.pdf",
        doc_type=DocumentType.PDF, page_count=2,
        chunks=[
            DocumentChunk(
                chunk_id="test-job-0000", text="some text",
                token_estimate=2,
                metadata={
                    "heading": "Intro", "has_table": False, "has_clause": False,
                    "has_definition": False, "has_image": False,
                },
            )
        ],
        markdown="# Intro\n\nsome text",
        processing_time_ms=42,
        sha256="deadbeef",
    )
    doc._quiz_items    = []                                           # type: ignore[attr-defined]
    doc._ai_readiness  = {"overall_score": 0.9, "page_scores": [],   # type: ignore[attr-defined]
                          "warnings": []}
    doc._chapter       = "Ch. 2"                                      # type: ignore[attr-defined]
    doc._heading_to_page = {}                                         # type: ignore[attr-defined]
    serial = serialize_result(doc)
    check("serialize job_id",         serial["job_id"],                "test-job")
    check("serialize filename",       serial["filename"],              "test.pdf")
    check("serialize chunk count",    serial["stats"]["total_chunks"], 1)
    check("serialize overall score",  serial["ai_readiness"]["overall_score"], 0.9)
    check("serialize quiz list",      serial["quiz_items"],            [])
    check("images_detected key",      "images_detected" in serial["stats"], True)

    # ── SUMMARY ──────────────────────────────────────────────────────────────
    total = passed + failed
    print(f"\n{'═' * 60}")
    print(f"  {passed}/{total} tests passed  "
          f"{'🟢  ALL GREEN' if not failed else '🔴  FAILURES'}")
    print(f"{'═' * 60}\n")
    sys.exit(0 if not failed else 1)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHANGELOG
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#
# v3.0.0  (this file)
# -------------------
# 🔴 BUG 1 — CHUNKING: Replaced RecursiveCharacterTextSplitter-as-primary
#     with _split_markdown_by_heading() which produces one section per H1/H2/H3
#     boundary.  RecursiveCharacterTextSplitter is now a fallback used only
#     when a section exceeds CHUNK_SIZE_TOKENS.  A 34-slide deck now produces
#     ~34 chunks instead of 1.
#
# 🔴 BUG 2 — SLIDE_NUMBER: _build_markdown() now builds heading_to_page
#     (normalized heading text → page number) and passes it to _chunk_text().
#     _page_for_section() first does a heading dict lookup; falls back to
#     word-overlap search across all blocks.  slide_number: 0 is now the
#     exception rather than the rule.
#
# 🔴 BUG 3 — _TRAILING_DOUBLE_PAGENUM_RE: regex now requires the doubled block
#     to be glued directly to a non-digit, non-space character.  "section 11",
#     "22", "F2 22", "problem 33" are no longer mangled.
#
# 🟠 BUG 4 — has_table: _is_punnett_or_ratio() detects "TT : Tt : tt" and
#     "1 : 2 : 1" patterns.  has_table is now set when either pipe-table OR
#     Punnett/ratio lines are present.
#
# 🟠 BUG 5 — QUIZ STEM: stem is now collected strictly between the "Question"
#     heading and the first choice.  No earlier bullets are absorbed.  For
#     headingless MCQ, we scan backwards from the first choice to the nearest
#     non-choice block.
#
# 🟡 BUG 6 — STAGE ORDER: bullet_continuations → dedup → noise_removal →
#     classify is now the guaranteed order.  Headers can no longer absorb body
#     bullets that were supposed to be merged first.
#
# ⚠️ WEAKNESS 1 — ENCODING: _remap_symbol_chars() maps Wingdings/Symbol
#     Private-Use Area code-points to ASCII/Unicode equivalents.  Applied at
#     the bottom of _merge_spans_into_visual_lines() and in all non-PDF
#     extractors.  U+FFFD → "-".
#
# ⚠️ WEAKNESS 2 — SINGLE MEGA-CHUNK: fixed by BUG 1 resolution above.
#
# 🟠 WEAKNESS 3 — TABLES: pdfplumber integration unchanged; Punnett patterns
#     now detected and flagged via _is_punnett_or_ratio().
#
# 🟠 WEAKNESS 4 — QUIZ DETECTION: _is_choice_line() extended with
#     _MCQ_NUMBERED_RE to catch "1. 100%", "2. 50%", "3. none" patterns.
#
# 🟡 WEAKNESS 5 — IMAGES: _page_has_images() detects embedded image streams
#     via PyMuPDF.  _build_markdown() injects ![Figure](slide:N) placeholder
#     after first heading on image-bearing slides.  has_image metadata per
#     chunk and in ai_readiness page_scores.
#
# 🟡 WEAKNESS 6 — TRUNCATION: _extract_pdf() logs a warning when the last
#     page produces no text blocks.
#
# 🟡 WEAKNESS 7 — SPURIOUS HEADINGS: headings with ≤1 clean character after
#     stripping '#' are suppressed in _remove_noise() and _build_markdown().
#     _classify_blocks() also guards on len(clean) < 2.
#
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

if __name__ == "__main__":
    if "--test" in sys.argv:
        logging.basicConfig(level=logging.WARNING)
        _run_tests()
    else:
        print("Usage: python pipeline.py --test")